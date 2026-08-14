"""
tests for detectors, replacements, document, and redactor
"""
import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from detectors import (
    _luhn,
    _detect_emails,
    _detect_phones,
    _detect_ssn,
    _detect_credit_cards,
    _detect_ips,
    _detect_dob,
    _resolve_overlaps,
    DetectedEntity,
)
from replacements import ReplacementMap


# --- luhn ---

def test_luhn_valid():
    assert _luhn("4532015112830366") is True


def test_luhn_invalid():
    assert _luhn("1234567890123456") is False


def test_luhn_too_short():
    assert _luhn("12345") is False


# --- email ---

def test_detect_email_basic():
    entities = _detect_emails("Contact us at alice@example.com for more info.")
    assert len(entities) == 1
    assert entities[0].type == "EMAIL"
    assert entities[0].text == "alice@example.com"


def test_detect_email_none():
    assert _detect_emails("No emails here.") == []


# --- phone ---

def test_detect_phone_us():
    entities = _detect_phones("Call (800) 555-1234 for support.")
    assert len(entities) >= 1
    assert entities[0].type == "PHONE"


# --- ssn ---

def test_detect_ssn_basic():
    entities = _detect_ssn("SSN: 123-45-6789")
    assert len(entities) == 1
    assert entities[0].type == "SSN"


def test_detect_ssn_invalid_prefix():
    # starts with 000 — should not match
    assert _detect_ssn("000-45-6789") == []


# --- credit card ---

def test_detect_cc_valid_luhn():
    # valid Visa test number
    entities = _detect_credit_cards("Card: 4532015112830366")
    assert len(entities) == 1
    assert entities[0].type == "CREDIT_CARD"


def test_detect_cc_invalid_luhn():
    assert _detect_credit_cards("Card: 1234567890123456") == []


# --- ip address ---

def test_detect_ip_public():
    entities = _detect_ips("Server at 203.0.113.42 is down.")
    assert len(entities) == 1
    assert entities[0].type == "IP_ADDRESS"


def test_detect_ip_loopback_skipped():
    # loopback should be skipped
    assert _detect_ips("Local: 127.0.0.1") == []


# --- date of birth ---

def test_detect_dob_with_context():
    entities = _detect_dob("Date of birth: 04/12/1985")
    assert len(entities) == 1
    assert entities[0].type == "DATE_OF_BIRTH"


def test_detect_dob_without_context():
    # date without DOB context should not be picked up by _detect_dob
    assert _detect_dob("The meeting is on 04/12/1985.") == []


# --- overlap resolution ---

def test_resolve_overlaps_keeps_higher_confidence():
    e1 = DetectedEntity("PERSON", "John Smith", 0, 10, 0.9)
    e2 = DetectedEntity("ORGANIZATION", "John Smith Corp", 0, 15, 0.5)
    result = _resolve_overlaps([e1, e2])
    assert len(result) == 1
    assert result[0].type == "PERSON"


def test_resolve_overlaps_non_overlapping():
    e1 = DetectedEntity("EMAIL", "a@b.com", 0, 7, 0.98)
    e2 = DetectedEntity("PERSON", "Alice", 20, 25, 0.9)
    result = _resolve_overlaps([e1, e2])
    assert len(result) == 2


# --- replacement map ---

def test_replacement_map_consistency():
    rmap = ReplacementMap()
    entity = DetectedEntity("EMAIL", "test@test.com", 0, 13, 0.98)
    first = rmap.get_or_create(entity)
    second = rmap.get_or_create(entity)
    assert first == second


def test_replacement_map_different_types():
    rmap = ReplacementMap()
    e1 = DetectedEntity("PERSON", "Alice", 0, 5, 0.9)
    e2 = DetectedEntity("ORGANIZATION", "Alice", 0, 5, 0.9)
    # same text but different type → can produce different values
    r1 = rmap.get_or_create(e1)
    r2 = rmap.get_or_create(e2)
    # both should be strings, just verify they exist
    assert isinstance(r1, str) and len(r1) > 0
    assert isinstance(r2, str) and len(r2) > 0


# --- redactor integration ---

def test_redact_produces_valid_docx():
    """smoke test: redact a trivial docx and open it back"""
    from docx import Document
    from redactor import redact

    # create a simple docx with known PII
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_in = tmp.name
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_out = tmp.name

    doc = Document()
    doc.add_paragraph("Contact alice@example.com or call (800) 555-1234.")
    doc.save(tmp_in)

    try:
        counts, entities = redact(tmp_in, tmp_out)
        # verify output is a valid docx
        out_doc = Document(tmp_out)
        full_text = " ".join(p.text for p in out_doc.paragraphs)
        # original email must be gone
        assert "alice@example.com" not in full_text
        assert isinstance(counts, dict)
    finally:
        os.unlink(tmp_in)
        try:
            os.unlink(tmp_out)
        except FileNotFoundError:
            pass
