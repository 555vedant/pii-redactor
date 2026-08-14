import copy
import re
from typing import List, Tuple

from docx import Document
from docx.oxml.ns import qn


def _iter_paragraphs(doc: Document):
    """yield all paragraphs including those inside table cells and text boxes"""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    # headers and footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr is not None:
                for para in hdr.paragraphs:
                    yield para


def read_paragraphs(path: str) -> Tuple[Document, List]:
    """return (doc, list of paragraph objects)"""
    doc = Document(path)
    return doc, list(_iter_paragraphs(doc))


def get_full_text(path: str) -> str:
    """join all paragraph text with newlines — used by detectors"""
    doc = Document(path)
    parts = []
    for para in _iter_paragraphs(doc):
        parts.append(para.text)
    return "\n".join(parts)


def _apply_replacements_to_para(para, replacements: List[Tuple[str, str]]):
    """replace text in a paragraph's runs, preserving run formatting.

    replacements is a list of (original, replacement) tuples, sorted
    longest-first to avoid partial matches.
    """
    if not replacements:
        return

    # build full paragraph text with run boundaries
    full_text = "".join(run.text for run in para.runs)
    if not any(orig in full_text for orig, _ in replacements):
        return

    # First try replacing within individual runs (preserves all formatting)
    for orig, repl in sorted(replacements, key=lambda x: -len(x[0])):
        for run in para.runs:
            if orig in run.text:
                run.text = run.text.replace(orig, repl)

    # Check if any replacements are still needed (meaning they spanned across runs)
    new_full_text = "".join(run.text for run in para.runs)
    spanned_replacements = [(orig, repl) for orig, repl in replacements if orig in new_full_text]
    
    if spanned_replacements:
        # Fallback for cross-run spans: collapse to first run
        modified = new_full_text
        for orig, repl in sorted(spanned_replacements, key=lambda x: -len(x[0])):
            modified = modified.replace(orig, repl)
        
        if modified != new_full_text and para.runs:
            para.runs[0].text = modified
            for run in para.runs[1:]:
                run.text = ""


def write_redacted_docx(original_path: str, text_replacements: List[Tuple[str, str]], output_path: str):
    """open a copy of the original docx, apply text replacements, save to output_path"""
    doc = Document(original_path)

    # sort by length descending so longer originals matched first
    sorted_repls = sorted(text_replacements, key=lambda x: -len(x[0]))

    for para in _iter_paragraphs(doc):
        _apply_replacements_to_para(para, sorted_repls)

    doc.save(output_path)
